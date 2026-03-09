from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

DARK_BLUE = RGBColor(0x1E,0x3A,0x5F)
MID_BLUE  = RGBColor(0x1D,0x4E,0xD8)
TEAL      = RGBColor(0x0D,0x94,0x88)
WHITE     = RGBColor(0xFF,0xFF,0xFF)
BLACK     = RGBColor(0x11,0x18,0x27)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(14); run.font.color.rgb = DARK_BLUE
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"6")
    bot.set(qn("w:space"),"1"); bot.set(qn("w:color"),"1D4ED8")
    pb.append(bot); pPr.append(pb)

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(11.5); run.font.color.rgb = MID_BLUE

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(10.5); run.font.color.rgb = TEAL

def body(text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.font.size = Pt(10); run.italic = italic; run.font.color.rgb = BLACK

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(0.25*(level+1))
    run = p.add_run(text)
    run.font.size = Pt(9.5); run.font.color.rgb = BLACK

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:fill"),"F3F4F6")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Courier New"; run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1F,0x29,0x37)

def info(label, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0,0)
    set_cell_bg(cell,"DBEAFE")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run(label+"  "); r1.bold=True; r1.font.size=Pt(9.5)
    r1.font.color.rgb = RGBColor(0x1E,0x40,0xAF)
    r2 = p.add_run(text); r2.font.size=Pt(9.5)
    r2.font.color.rgb = RGBColor(0x1E,0x3A,0x5F)
    doc.add_paragraph()

def table(headers, rows, widths=None):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hr = tbl.rows[0]
    for i,h in enumerate(headers):
        c = hr.cells[i]; set_cell_bg(c,"1D4ED8")
        p = c.paragraphs[0]
        run = p.add_run(h); run.bold=True; run.font.size=Pt(9)
        run.font.color.rgb=WHITE; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for ri,row in enumerate(rows):
        tr = tbl.rows[ri+1]
        bg = "F8FAFC" if ri%2==0 else "FFFFFF"
        for ci,val in enumerate(row):
            c=tr.cells[ci]; set_cell_bg(c,bg)
            p=c.paragraphs[0]; run=p.add_run(str(val))
            run.font.size=Pt(8.5); run.font.color.rgb=BLACK
    if widths:
        for i,w in enumerate(widths):
            for row in tbl.rows:
                row.cells[i].width=Inches(w)
    doc.add_paragraph()

# ── TITLE ─────────────────────────────────────────────────────────────────────
tb = doc.add_table(rows=1,cols=1); tb.alignment=WD_TABLE_ALIGNMENT.LEFT
tc = tb.cell(0,0); set_cell_bg(tc,"1E3A5F")
tp = tc.paragraphs[0]; tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before=Pt(16); tp.paragraph_format.space_after=Pt(16)
r=tp.add_run("German Sign Language (DGS)  Text-to-Sign Pipeline")
r.bold=True; r.font.size=Pt(19); r.font.color.rgb=WHITE
doc.add_paragraph()
ps=doc.add_paragraph(); ps.alignment=WD_ALIGN_PARAGRAPH.CENTER
rs=ps.add_run("NLP Pipeline Technical Report  \u00b7  NLIP Lab \u00b7  Semester 6")
rs.font.size=Pt(11); rs.italic=True; rs.font.color.rgb=RGBColor(0x6B,0x72,0x80)
doc.add_paragraph()

# ── 1. OVERVIEW ───────────────────────────────────────────────────────────────
h1("1.  Project Overview")
body(
    "This project builds an end-to-end pipeline that converts German text sentences into "
    "Deutsche Gebaerdensprache (DGS) animations. It combines four NLP lab exercises with "
    "computer-vision keypoint extraction and a real-time WebSocket recognition service. "
    "All NLP processing is implemented from scratch in Python without pre-trained model "
    "dependencies, applied to real DGS-Korpus data."
)
info("Dataset:",
    "DGS-Korpus (Institut fuer Deutsche Gebaerdensprache, Universitaet Hamburg). "
    "ELAN EAF annotation files time-align gloss-level transcriptions to signer video. "
    "Training uses Signer A right-hand gloss tier (Lexem_Gebaerde_r_A) from one reference EAF.")

h2("Pipeline Stages at a Glance")
table(
    ["#","Stage","Module / File","NLP Relevance"],
    [
        ["1","Data Ingestion",     "ingest_conversation.py",           "EAF XML parse \u2192 German sentences + gloss sequences"],
        ["2","Keypoint Extraction","extract_keypoints_for_segments.py","MediaPipe landmarks \u2192 134-D keypoint arrays (.npz)"],
        ["3","Gloss Dictionary",   "build_gloss_dictionary.py",        "EAF gloss tier \u2192 canonical per-sign clip index"],
        ["4","Text Normalisation", "nlp_pipeline.py  (Lab 1)",         "Tokenization, stop-word removal, stemming, lemmatization"],
        ["5","Feature Engineering","nlp_pipeline.py  (FE Exercise)",   "BoW, TF-IDF, N-grams, PMI/PPMI, word embeddings"],
        ["6","Language Modelling", "nlp_pipeline.py  (Exercise 3)",    "N-gram LM with Laplace smoothing; perplexity; re-ranking"],
        ["7","POS Tagging",        "nlp_pipeline.py  (Exercise 4)",    "HMM + Viterbi decoding on German input text"],
        ["8","Translation",        "text_to_gloss_map.py",             "German \u2192 ordered DGS gloss sequence"],
        ["9","Live Recognition",   "main.py  /ws/live_recognition",    "MediaPipe hand detection + cosine centroid matching"],
        ["10","API & Frontend",    "main.py (FastAPI) + React/Vite",   "REST + WebSocket; 3-D skeleton animation viewer"],
    ],
    widths=[0.25,1.45,2.0,2.8]
)

# ── 2. DATA INGESTION ─────────────────────────────────────────────────────────
h1("2.  Data Ingestion  (ingest_conversation.py)")
body(
    "ELAN Annotation Format (EAF) files are XML documents. Each file can hold multiple "
    "parallel tiers aligned to one or more signer videos. The ingestion script walks every "
    "EAF in the dataset directory, extracts sentence-level German text and the corresponding "
    "right-hand gloss sequence, and writes a unified segments_manifest.json."
)
h2("EAF Tiers Consumed")
table(
    ["Tier ID","Content","Usage"],
    [
        ["Deutsche_Uebersetzung_A",  "Sentence-level German translation",           "Source text for NLP pipeline"],
        ["Lexem_Gebaerde_r_A",        "Right-hand gloss label per sign (Signer A)", "Training gloss vocabulary"],
        ["TIME_ORDER / TIME_SLOT",    "Absolute millisecond offsets",               "Segment boundary alignment"],
    ],
    widths=[2.2,2.4,2.0]
)
h2("Output: segments_manifest.json")
code(
    '{ "segments": [\n'
    '    { "id": "seg_001",\n'
    '      "german_text": "Wie mein Leben aussieht?",\n'
    '      "gloss_sequence": ["SEHEN1*","SELBST1A*","LEBEN1A*","SEHEN1"],\n'
    '      "start_ms": 240,  "end_ms": 2160 },\n'
    '    ...\n'
    '] }'
)
body(
    "This manifest is the single source-of-truth consumed by the NLP language model, "
    "the HMM tagger training loop, and the feature engineering corpus builder."
)

# ── 3. KEYPOINT EXTRACTION ────────────────────────────────────────────────────
h1("3.  Keypoint Extraction  (extract_keypoints_for_segments.py)")
body(
    "Each video segment is processed frame-by-frame with two Google MediaPipe landmark models. "
    "Results are stored as compressed NumPy archives consumed by both the animation renderer "
    "and the gloss dictionary builder."
)
table(
    ["Model","Joints detected","Coordinates","Contribution to 134-D vector"],
    [
        ["pose_landmarker_lite.task","25 body joints","x, y (normalised 0-1)","50 values  (cols 0-49)"],
        ["hand_landmarker.task","21 joints \u00d7 2 hands","x, y (normalised 0-1)","84 values  (cols 50-133)"],
    ],
    widths=[2.0,1.55,1.7,2.0]
)
body(
    "The hand columns (50:134) carry the highest discriminative power for gloss recognition; "
    "the body columns anchor the spatial reference frame for animation."
)

# ── 4. GLOSS DICTIONARY ───────────────────────────────────────────────────────
h1("4.  Gloss Dictionary  (build_gloss_dictionary.py)")
body(
    "The gloss dictionary maps every unique DGS gloss label to a keypoint clip sliced "
    "directly from the training motion data. It functions as the sign lexicon for both "
    "the animation renderer and the live recogniser."
)
h2("Build Algorithm")
bullet("Parse EAF TIME_SLOT table \u2192 {slot_id: milliseconds} map.")
bullet("Iterate Lexem_Gebaerde_r_A; per annotation extract (start_ms, end_ms, gloss_label).")
bullet("Find segment .npz whose window contains the gloss (exact containment; best-overlap fallback).")
bullet("Convert ms to frame indices:  f = round((t_ms \u2212 seg_start_ms) \u00d7 fps / 1000)")
bullet("Slice  kp[f1:f2, :]  and save as  gloss_clips/<SAFE_NAME>_<n>.npz.")
bullet("Record first occurrence per unique gloss in gloss_dictionary.json.")
info("Scale:",
    "Training EAF contains 741 right-hand gloss annotations. After filtering clips "
    "shorter than 2 frames, ~700 unique gloss entries are indexed. "
    "One canonical centroid per gloss is used for live recognition.")

# ── 5. NLP PIPELINE ───────────────────────────────────────────────────────────
h1("5.  NLP Pipeline  (nlp_pipeline.py)  \u2014 Primary Focus")
body(
    "The NLP module satisfies four progressive lab exercises applied directly to the "
    "German-to-gloss translation task. All four components are accessible via dedicated "
    "REST endpoints and rendered in the frontend NLP Analysis panel."
)

# 5.1 LAB 1
h2("5.1  Lab 1 \u2014 Tokenization & Text Normalisation")
body(
    "German text passes through six independent stages before any statistical processing. "
    "Each stage is independently accessible and its output is returned by /api/nlp/analyze."
)
table(
    ["Stage","Function","Implementation","Example output"],
    [
        ["Word tokenisation",     "word_tokenize()",        "Regex  [a-zA-Z\u00e4\u00f6\u00fc\u00c4\u00d6\u00dc\u00df]+",   '"Ich bin taub" \u2192 ["Ich","bin","taub"]'],
        ["Sentence tokenisation", "sentence_tokenize()",    "NLTK sent_tokenize (German) \u2192 regex fallback","Splits on .!?"],
        ["Character tokenisation","char_tokenize()",        "list(text)","'DGS' \u2192 ['D','G','S']"],
        ["Subword BPE",           "subword_tokenize_bpe()", "Custom BPE (Sennrich et al. 2016)","aufgewachsen \u2192 [auf,ge,wachs,en]"],
        ["Stop-word removal",     "remove_stopwords()",     "Custom German set (~80 tokens)","drops: ich,und,der,die,das,..."],
        ["Stemming",              "stem_tokens()",          "NLTK SnowballStemmer (german)","leben\u2192leb, gehörlos\u2192gehörlos"],
        ["Lemmatisation",         "lemmatize_tokens()",     "Rule-based suffix stripping","aufgewachsen\u2192aufgewachs"],
    ],
    widths=[1.55,1.55,2.0,1.65]
)

h3("Byte-Pair Encoding (BPE) \u2014 Custom Implementation")
body(
    "The BPE tokeniser follows Sennrich et al. (2016). Words are split into characters "
    "with a </w> end-of-word marker. The most frequent adjacent character pair is merged "
    "iteratively until a vocab_size budget is exhausted, then the learned merges are "
    "applied to segment the input text into subword units."
)
code(
    "Input word: 'gehörlos'\n"
    "Initial   : g e h ö r l o s </w>\n"
    "Merge 1   : ge h ö r l o s </w>   (ge most frequent)\n"
    "Merge 2   : ge hö r l o s </w>\n"
    "...                              \n"
    "Final BPE : ['ge','hö','rlos']"
)

h3("Rule-Based Lemmatisation")
body(
    "Inflectional suffixes are stripped in priority order. A minimum stem length of 3 "
    "characters is enforced so that stripping does not destroy short words. "
    "Twenty-six suffix rules cover the most productive German morphological patterns "
    "(e.g. -ungen\u2192-ung, -ieren\u2192-ieren, -enden\u2192-en, -ten\u2192-t)."
)

# 5.2 FEATURE ENGINEERING
h2("5.2  Feature Engineering \u2014 BoW, TF-IDF, N-grams, PMI, Embeddings")
body(
    "The feature engineering module builds numerical representations of the German corpus "
    "(first 50 manifest segments by default) to support statistical analysis and "
    "translation re-ranking."
)

h3("Bag-of-Words")
body(
    "build_bow() constructs a dense document-term count matrix. Vocabulary is sorted "
    "alphabetically. Sparsity (fraction of zero cells) is reported as a corpus health metric."
)
code(
    "Matrix[i][j] = count of vocab_term[j] in segment sentence i\n"
    "Sparsity     = |zero cells| / (|docs| x |vocab|)"
)

h3("N-gram Features")
body(
    "build_ngrams(tokens, n) yields all contiguous n-tuples using a sliding window. "
    "Unigram, bigram, and trigram vocabulary sizes illustrate the feature-space explosion:"
)
code(
    "Tokens : ['taub','gehörlos','aufgewachsen']\n"
    "Bigrams: [('taub','gehörlos'), ('gehörlos','aufgewachsen')]\n"
    "Vocab grows exponentially:  |V_bigram| >> |V_unigram|"
)

h3("TF-IDF")
body(
    "compute_tfidf() uses sklearn TfidfVectorizer with a German-aware token pattern. "
    "Manual fallback (no sklearn) implements:"
)
code(
    "TF(t,d)   = count(t,d) / |d|\n"
    "IDF(t)    = log( (N+1) / (df(t)+1) ) + 1    [smooth IDF]\n"
    "TFIDF     = TF x IDF\n"
    "Top-10 terms by corpus-average score surfaced via API."
)

h3("Pointwise Mutual Information (PMI / PPMI)")
body(
    "compute_pmi() measures association between word pairs within a sliding context window "
    "(default width \u00b12 tokens). PPMI clamps negative values to zero, suppressing "
    "accidental anti-correlations from sparse counts."
)
code(
    "PMI(w1,w2)  = log2[ P(w1,w2) / (P(w1) x P(w2)) ]\n"
    "PPMI(w1,w2) = max(0,  PMI(w1,w2))\n"
    "P(w1,w2) estimated from sliding-window co-occurrence counts."
)

h3("Word Embeddings")
body(
    "Sentence-transformers (multilingual-MiniLM-L6-v2, 384 dimensions) provide dense "
    "vector representations used as a fallback similarity measure when exact or stem-based "
    "gloss matches fail in the translation mapper."
)

# 5.3 N-GRAM LM
h2("5.3  Exercise 3 \u2014 N-gram Language Model for Gloss Sequences")
body(
    "Three NGramLM instances (n = 1, 2, 3) are trained over the gloss sequences in the "
    "manifest. The key insight is that DGS gloss order follows its own sequential "
    "regularities distinct from German word order; modelling these with an n-gram LM "
    "enables data-driven re-ranking of translation candidates."
)

h3("NGramLM Class")
table(
    ["Method","Purpose","Formula / Detail"],
    [
        ["train(sequences)",         "Count n-grams from padded sequences","Pads with <s>x(n-1) and </s>; increments count and context dicts"],
        ["log_prob(ngram)",          "Laplace-smoothed log-probability",   "log[(c(ngram)+1) / (c(context)+|V|)]"],
        ["score_sequence(seq)",      "Total log-prob of a gloss sequence", "Sum log_prob over all sliding n-gram windows"],
        ["perplexity(test_seqs)",    "Model quality on held-out data",     "exp(\u2212(1/N) \u03a3 log P(w_i|ctx))  \u2014  lower is better"],
        ["score_candidates(cands)",  "Rank candidate gloss sequences",     "Returns sorted list (seq, log_prob, avg_per_token)"],
        ["get_counts_table(k)",      "Top-k n-gram counts for API display","Sorted descending by frequency"],
    ],
    widths=[1.6,1.8,3.0]
)

h3("Laplace (Add-One) Smoothing")
body(
    "Every unseen n-gram is assigned a pseudo-count of 1, preventing zero probability "
    "for glosses absent from training data -- critical given the limited corpus size."
)
code(
    "P_Laplace(w_n | w_{n-1}...w_1)\n"
    "  =  ( c(w_1...w_n) + 1 )  /  ( c(w_1...w_{n-1}) + |V| )\n\n"
    "|V| includes the special <s> and </s> boundary tokens."
)

h3("Translation Candidate Ranking")
body(
    "GlossLanguageModel.select_best_translation() receives a list of candidate gloss "
    "sequences (produced by different stop-word or ordering strategies), scores each "
    "with the bigram model, and returns the highest log-probability sequence. "
    "This acts as a lightweight fluency re-ranker without any additional training."
)

# 5.4 HMM
h2("5.4  Exercise 4 \u2014 HMM POS Tagger with Viterbi Decoding")
body(
    "An HMM POS tagger is trained on the German sentences from the manifest using "
    "auto-generated heuristic labels. It performs Viterbi decoding to find the most "
    "probable POS tag sequence, which then guides which tokens the gloss mapper should "
    "look up versus discard."
)

h3("HMM Parameters")
table(
    ["Parameter","Symbol","Estimation","Smoothing"],
    [
        ["Initial probability","pi(t)",    "Fraction of sentences where tag t is first",   "Add-1 over |tags|"],
        ["Transition prob.",   "P(t|t-1)", "Bigram tag-pair counts from training",          "Add-1 over |tags|"],
        ["Emission prob.",     "P(w|t)",   "Per-tag word frequency from training sentences","Add-1 over |vocab(t)|"],
    ],
    widths=[1.6,0.9,2.4,1.5]
)

h3("Tag Set  (11 Universal POS Tags)")
body("NOUN  |  VERB  |  ADJ  |  ADV  |  PRON  |  DET  |  ADP  |  CONJ  |  NUM  |  PUNCT  |  X")

h3("Heuristic Auto-Labelling for Training")
body(
    "Since no annotated German treebank is bundled, training labels are produced by "
    "priority-ordered regex patterns. Selected rules:"
)
table(
    ["Pattern","Assigned Tag","Rationale"],
    [
        ["^[A-Z\u00c4\u00d6\u00dc][a-z\u00e4\u00f6\u00fc]+$  (mid-sentence capital)","NOUN","German nouns are capitalised"],
        ["ist|bin|bist|sind|war|hat|haben|werden","VERB","Common German copula / auxiliary forms"],
        ["ich|du|er|sie|es|wir|ihr","PRON","Personal pronoun list"],
        ["der|die|das|ein|eine|einem|einer","DET","German articles"],
        ["auf|in|an|bei|von|mit|fuer|ueber|zu","ADP","Common prepositions"],
        ["und|oder|aber|weil|dass|wenn","CONJ","Coordinating / subordinating conjunctions"],
        ["suffix -ung/-heit/-keit","NOUN","Nominalisation suffixes"],
        ["suffix -lich/-isch/-haft","ADJ","Adjectival derivational suffixes"],
        ["suffix -en/-eln/-ieren","VERB","Infinitive / verbal suffixes"],
    ],
    widths=[2.6,1.0,2.8]
)

h3("Viterbi Decoding")
body(
    "For an input of N tokens and T=11 tags the algorithm fills an N\u00d7T matrix in "
    "O(N\u00b7T\u00b2) time, then backtracks to recover the best tag sequence."
)
code(
    "--- Initialisation (position t=0) ---\n"
    "V[0][i] = log pi(tag_i) + log P(token_0 | tag_i)\n\n"
    "--- Recursion (t = 1 ... N-1) ---\n"
    "V[t][i] = max_j [ V[t-1][j] + log P(tag_i|tag_j) ]\n"
    "        + log P(token_t | tag_i)\n"
    "BP[t][i]= argmax_j above\n\n"
    "--- Backtrack ---\n"
    "best[N-1] = argmax_i V[N-1][i]\n"
    "best[t]   = BP[t+1][best[t+1]]   for t = N-2 ... 0"
)
body(
    "The complete Viterbi log-probability matrix ([tags x tokens]) is returned by "
    "get_viterbi_matrix_display() and rendered as a scrollable heatmap table in the "
    "frontend NLP Analysis panel."
)

h3("Role in the Translation Pipeline")
body(
    "POS labels assigned by the HMM determine gloss lookup eligibility. "
    "Tokens tagged as NOUN, VERB, ADJ, and content-ADV are mapped to gloss dictionary entries. "
    "Tags DET, ADP, CONJ, PRON, and PUNCT are dropped -- German function words have no "
    "direct DGS equivalent in the corpus annotation scheme."
)

# ── 6. TRANSLATION ────────────────────────────────────────────────────────────
h1("6.  Text-to-Gloss Translation  (text_to_gloss_map.py)")
body(
    "GlossMapper converts a German sentence to an ordered DGS gloss sequence. "
    "Two modes are exposed via the /api/translate endpoint."
)
table(
    ["Mode","API flag","Description"],
    [
        ["Standard","chained=false",
         "Normalise \u2192 remove stop-words \u2192 HMM POS filter \u2192 gloss dictionary lookup per content word. "
         "Returns one gloss per matched token."],
        ["Chained","chained=true",
         "Each content word is matched to its best segment clip. Clips are concatenated with "
         "inter-sign transition frames producing smoother animations."],
    ],
    widths=[1.0,1.1,4.3]
)
h2("Lookup Cascade")
bullet("1. Exact match   \u2014 gloss_dictionary key == normalised token.")
bullet("2. Stem match    \u2014 SnowballStemmer applied to both token and all dictionary keys.")
bullet("3. Embedding sim \u2014 cosine similarity of sentence-transformer vectors (OOV fallback).")
bullet("4. LM re-ranking \u2014 bigram GlossLM scores all candidate orderings; highest wins.")

h2("Animation Assembly")
body(
    "POST /api/motion/by_glosses loads each gloss clip .npz, concatenates with transitions, "
    "and streams the combined keypoint array to the 3-D skeleton viewer."
)
code(
    "Between clip_A [T1,134] and clip_B [T2,134]:\n"
    "  Hold  : repeat last frame of clip_A for 3 frames   (~120 ms @ 25 fps)\n"
    "  Lerp  : alpha*last_A + (1-alpha)*first_B for 5 frames  (~200 ms)\n"
    "Result  : smooth coarticulation transition between signs"
)

# ── 7. LIVE RECOGNITION ───────────────────────────────────────────────────────
h1("7.  Live Sign Recognition  (/ws/live_recognition)")
body(
    "A WebSocket endpoint provides real-time DGS recognition from a webcam or uploaded "
    "video file. Frames are streamed as base64 JPEG; the server runs MediaPipe hand "
    "detection and performs zero-shot nearest-centroid matching against the gloss dictionary."
)
h2("Per-Frame Recognition Pipeline")
table(
    ["Step","Detail"],
    [
        ["Receive",       'Client sends JSON  {"frame": "<base64-JPEG>"}  every 120 ms.'],
        ["Decode",        "base64 \u2192 bytes \u2192 cv2.imdecode \u2192 RGB NumPy array [H,W,3]."],
        ["Hand detect",   "MediaPipe HandLandmarker (IMAGE mode) \u2192 up to 2 hands, 21 landmarks each."],
        ["Feature vector","left-hand 21\u00d72 + right-hand 21\u00d72 concatenated \u2192 84-D query vector q."],
        ["Cosine match",  "sim(q,c_i) = (q\u00b7c_i) / (||q||\u00b7||c_i||)  for every centroid c_i."],
        ["Top-3 output",  "Sorted descending \u2192 best gloss + confidence + top-3 list sent to client."],
        ["Auto-append",   "Frontend appends to history strip only when confidence > 60%."],
    ],
    widths=[1.1,5.3]
)
h2("Centroid Computation  (_load_gloss_centroids)")
code(
    "kp  [T, 134]  loaded from gloss_clips/<gloss>.npz\n"
    "hand_vec = kp[:, 50:].mean(axis=0)   # shape [84] - time-averaged hand landmarks\n"
    "centroids[gloss_label] = hand_vec"
)
h2("Frame Capture Parameters")
bullet("Capture interval : 120 ms  \u2248  8 fps")
bullet("Canvas resolution : 320 \u00d7 240 px  (downscaled from 640\u00d7480 webcam)")
bullet("JPEG quality : 60%  \u2014  balances bandwidth vs landmark accuracy")
bullet("MediaPipe mode : IMAGE (no temporal smoothing, stateless per frame)")

# ── 8. REST API ───────────────────────────────────────────────────────────────
h1("8.  REST API  (FastAPI / main.py)")
table(
    ["Endpoint","Method","Purpose"],
    [
        ["/api/health",              "GET",       "Liveness probe \u2192 {status: ok}"],
        ["/api/translate",           "POST",      "German text \u2192 gloss list (standard or chained)"],
        ["/api/motion/by_glosses",   "POST",      "Gloss list \u2192 concatenated keypoint animation"],
        ["/api/motion/chained",      "POST",      "Segment IDs \u2192 chained keypoint animation"],
        ["/api/motion/{segment_id}", "GET",       "Single segment keypoints"],
        ["/api/nlp/analyze",         "POST",      "Full NLP pipeline on input German text"],
        ["/api/nlp/gloss_lm",        "GET",       "N-gram LM stats and top n-gram frequency tables"],
        ["/api/nlp/score_glosses",   "POST",      "Score gloss sequence with all three LM orders"],
        ["/api/nlp/pos_tables",      "GET",       "HMM transition and emission probability tables"],
        ["/ws/live_recognition",     "WebSocket", "Real-time per-frame gloss recognition stream"],
    ],
    widths=[2.2,0.9,3.3]
)

# ── 9. FRONTEND ───────────────────────────────────────────────────────────────
h1("9.  Frontend  (React / Vite)  \u2014 Brief")
table(
    ["Component","Role"],
    [
        ["App.jsx",               "Root layout; two-column grid; panel routing."],
        ["NLPAnalysisPanel.jsx",  "Text input \u2192 /api/nlp/analyze \u2192 renders tokenisation, POS table, Viterbi matrix, TF-IDF, PMI tabs."],
        ["SkeletonViewer3D.jsx",  "Renders 134-D keypoint array as animated 3-D stick figure (Three.js / React-Three-Fiber)."],
        ["WebcamRecognition.jsx", "Live camera + video-file modes \u2192 WebSocket frame streaming \u2192 gloss chip history."],
    ],
    widths=[2.0,4.4]
)
body("Stack: React 18, Vite 5, Three.js, React-Three-Fiber. No UI component library.", italic=True)

# ── 10. END-TO-END FLOW ───────────────────────────────────────────────────────
h1("10.  End-to-End Data Flow  (Text-to-Sign Path)")
code(
    'Input  "Ich bin taub aufgewachsen."\n'
    '   |\n'
    '   |  word_tokenize()       ["Ich", "bin", "taub", "aufgewachsen"]\n'
    '   |  remove_stopwords()    ["taub", "aufgewachsen"]\n'
    '   |  stem_tokens()         ["taub", "aufgewachs"]\n'
    '   |  HMM viterbi()         [("taub","ADJ"), ("aufgewachsen","VERB")]\n'
    '   |\n'
    '   |  GlossMapper.translate()\n'
    '   |    exact  "taub"       -> TAUB-GEHÖRLOS1A\n'
    '   |    stem   "aufgewachs" -> AUFWACHSEN1A*\n'
    '   |\n'
    '   |  GlossLM.select_best_translation()\n'
    '   |    bigram score ["TAUB-GEHÖRLOS1A","AUFWACHSEN1A*"] = -4.21  (ranked #1)\n'
    '   |\n'
    '   |  POST /api/motion/by_glosses\n'
    '   |    load TAUB-GEHÖRLOS1A_0.npz  -> [T1, 134]\n'
    '   |    load AUFWACHSEN1A_star_0.npz-> [T2, 134]\n'
    '   |    _concat_with_transitions()  -> [T_total, 134]\n'
    '   |\n'
    '   v\n'
    '  SkeletonViewer3D renders animated DGS sign sequence'
)

# ── 11. DESIGN DECISIONS ──────────────────────────────────────────────────────
h1("11.  Key Design Decisions & Limitations")
table(
    ["Aspect","Decision","Limitation / Trade-off"],
    [
        ["Gloss recognition",
         "Zero-shot cosine similarity on time-averaged hand centroids. No classifier training.",
         "Sensitive to signer variation; centroids from a single reference signer."],
        ["LM order",
         "Bigram used for re-ranking (trigram too sparse on ~700 training sequences).",
         "High perplexity on OOV gloss sequences; limited generalisation."],
        ["POS labelling",
         "Regex heuristics; no external treebank required.",
         "Lower accuracy than spaCy German model; labelling errors propagate downstream."],
        ["Smoothing",
         "Laplace add-one: simple, always non-zero, no held-out data needed.",
         "Over-smooths when |V| is large relative to corpus; Good-Turing would be better."],
        ["Gloss variants",
         "Dictionary stores first occurrence only (one centroid per unique label string).",
         "SEHEN1 vs SEHEN1* (phonological variant) collapse to one centroid."],
        ["Signer scope",
         "Training on Signer A right-hand tier only.",
         "Live recognition degrades when Signer B dominates the camera frame."],
        ["BPE vocab",
         "vocab_size=50 chosen to suit short German sign-language sentences.",
         "Optimal vocab size not tuned; may under- or over-segment for long texts."],
    ],
    widths=[1.4,2.45,2.55]
)

# SAVE
out = r"d:\College\Sem 6\Lab\NLIP\Project\signlanguage-german-text2sign\DGS_Pipeline_Report.docx"
doc.save(out)
print("Saved ->", out)
